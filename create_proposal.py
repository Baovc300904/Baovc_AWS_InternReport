from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Add header and footer with page numbers
section = doc.sections[0]

# Header
header = section.header
header_para = header.paragraphs[0]
header_para.text = "AWS First Cloud AI Journey – Project Plan"
header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
header_para.runs[0].font.size = Pt(9)
header_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# Footer with page numbers
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add page number field
run = footer_para.add_run()
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# Create page number field
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "PAGE"

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# Cover Page
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run('AWS FIRST CLOUD AI JOURNEY – PROJECT PLAN')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = subtitle.add_run('Devteria Team – First Cloud Journey – Devteria Game Store Platform')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = date_para.add_run('December 8, 2025')
run.font.size = Pt(12)

doc.add_page_break()

# Table of Contents
toc_heading = doc.add_heading('TABLE OF CONTENTS', 0)
toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Create TOC table for better alignment
toc_table = doc.add_table(rows=18, cols=2)
toc_table.style = 'Table Grid'
toc_table.autofit = False
toc_table.allow_autofit = False

# Set column widths
for row in toc_table.rows:
    row.cells[0].width = Inches(5.5)
    row.cells[1].width = Inches(1.0)

toc_items = [
    ('1', 'BACKGROUND AND MOTIVATION', '3'),
    ('     1.1', 'EXECUTIVE SUMMARY', '3'),
    ('     1.2', 'PROJECT SUCCESS CRITERIA', '4'),
    ('     1.3', 'ASSUMPTIONS', '4'),
    ('2', 'SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM', '5'),
    ('     2.1', 'TECHNICAL ARCHITECTURE DIAGRAM', '5'),
    ('     2.2', 'TECHNICAL PLAN', '6'),
    ('     2.3', 'PROJECT PLAN', '7'),
    ('     2.4', 'SECURITY CONSIDERATIONS', '8'),
    ('3', 'ACTIVITIES AND DELIVERABLES', '9'),
    ('     3.1', 'ACTIVITIES AND DELIVERABLES', '9'),
    ('     3.2', 'OUT OF SCOPE', '10'),
    ('     3.3', 'PATH TO PRODUCTION', '10'),
    ('4', 'EXPECTED AWS COST BREAKDOWN BY SERVICES', '11'),
    ('5', 'TEAM', '12'),
    ('6', 'RESOURCES & COST ESTIMATES', '13'),
    ('7', 'ACCEPTANCE', '14'),
]

for i, (num, title, page) in enumerate(toc_items):
    cell_left = toc_table.rows[i].cells[0]
    cell_right = toc_table.rows[i].cells[1]
    
    # Left cell: number and title
    cell_left.text = f'{num}     {title}'
    cell_left.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Right cell: page number (right-aligned)
    cell_right.text = page
    cell_right.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell_right.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Remove borders
    for border_name in ['top', 'bottom', 'start', 'end', 'insideH', 'insideV']:
        element = cell_left._element.get_or_add_tcPr()
        tcBorders = element.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = element.makeelement(qn('w:tcBorders'))
            element.append(tcBorders)
        border = tcBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = tcBorders.makeelement(qn(f'w:{border_name}'))
            tcBorders.append(border)
        border.set(qn('w:val'), 'none')
        
        element = cell_right._element.get_or_add_tcPr()
        tcBorders = element.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = element.makeelement(qn('w:tcBorders'))
            element.append(tcBorders)
        border = tcBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = tcBorders.makeelement(qn(f'w:{border_name}'))
            tcBorders.append(border)
        border.set(qn('w:val'), 'none')

doc.add_page_break()

# Section 1: BACKGROUND AND MOTIVATION
doc.add_heading('1  BACKGROUND AND MOTIVATION', 1)

doc.add_heading('1.1  EXECUTIVE SUMMARY', 2)
doc.add_paragraph(
    'Devteria Game Store is a scalable cloud-based e-commerce platform for digital game licensing and distribution. '
    'Built entirely on AWS infrastructure, it delivers secure user authentication, real-time inventory management, '
    'automated order processing, and global content delivery.'
)

doc.add_paragraph(
    'The platform addresses critical challenges faced by traditional digital stores:'
)
challenges = [
    'Traffic spike handling during sales events',
    'Complex authentication processes reducing conversions',
    'Manual inventory management causing overselling',
    'Lack of real-time analytics and insights',
    'High infrastructure costs for peak capacity planning'
]
for challenge in challenges:
    doc.add_paragraph(challenge, style='List Bullet')

doc.add_paragraph('Devteria leverages AWS services including:')
services = [
    'CloudFront + S3 for fast global content delivery',
    'Cognito for secure authentication with MFA',
    'API Gateway + Lambda for serverless backend logic',
    'RDS PostgreSQL + S3 for reliable data storage',
    'SQS + SNS for asynchronous order processing',
    'CodePipeline for continuous deployment'
]
for service in services:
    doc.add_paragraph(service, style='List Bullet')

doc.add_paragraph(
    'The solution supports thousands of concurrent users with 99.9% uptime guarantee and achieves cost efficiency '
    'through serverless architecture that scales automatically based on demand.'
)

# 1.2 PROJECT SUCCESS CRITERIA
doc.add_heading('1.2  PROJECT SUCCESS CRITERIA', 2)
criteria = [
    'Achieve page load time of less than 2 seconds globally through CloudFront CDN',
    'Maintain 99.9% system uptime using AWS managed services and Multi-AZ deployment',
    'Handle 10x traffic spikes automatically through Lambda auto-scaling and EC2 Auto Scaling Groups',
    'Ensure zero security breaches with Cognito MFA, IAM policies, and encryption at rest/transit',
    'Complete CI/CD pipeline enabling deployments in less than 10 minutes via CodePipeline',
    'Reduce cart abandonment by 40% through simplified Cognito authentication flow',
    'Decrease infrastructure management time by 60% using serverless components',
    'Maintain stable performance for 10,000 monthly active users with room to scale to 100K+',
    'Ensure database consistency with RDS automated backups and point-in-time recovery (7-day retention)',
    'Achieve PCI-ready security posture for payment processing integration'
]
for criterion in criteria:
    doc.add_paragraph(criterion, style='List Bullet')

# 1.3 ASSUMPTIONS
doc.add_heading('1.3  ASSUMPTIONS', 2)
assumptions = [
    'The customer will provide accurate business requirements and game licensing workflows',
    'GitLab repository and credentials will be accessible for CI/CD integration',
    'All workloads will operate within one AWS Region (us-east-1 or ap-southeast-1)',
    'No on-premise system integration required during the initial phase',
    'Frontend code (React/Next.js) supports static hosting via S3 + CloudFront',
    'The project will not include Machine Learning or AI recommendation modules in Phase 1',
    'Traffic is expected to remain within reasonable limits (10,000 users initially)',
    'Payment gateway integration will use third-party APIs (Stripe/PayPal) via Lambda',
    'Game files will be uploaded manually or via admin dashboard to S3',
    'Customer accepts AWS Free Tier limitations and cost estimates provided'
]
for assumption in assumptions:
    doc.add_paragraph(assumption, style='List Bullet')

doc.add_page_break()

# Section 2: SOLUTION ARCHITECTURE
doc.add_heading('2  SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM', 1)

doc.add_heading('2.1  TECHNICAL ARCHITECTURE DIAGRAM', 2)

# Add architecture diagram if it exists
diagram_path = r'static\images\5-Workshop\architecture.png'
if os.path.exists(diagram_path):
    doc.add_picture(diagram_path, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Add spacing after image
else:
    # Add placeholder text if diagram doesn't exist
    p = doc.add_paragraph()
    p.add_run('[Architecture Diagram - See technical documentation for visual representation]').italic = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

doc.add_paragraph('The Devteria Game Store architecture consists of the following layers:')

doc.add_paragraph().add_run('1. Edge Networking Layer').bold = True
edge_items = [
    'Route 53: DNS management and domain hosting',
    'CloudFront: Global CDN for caching static assets and reducing latency (<2s page load)',
    'AWS Shield: DDoS protection (Standard tier included)'
]
for item in edge_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('2. Application Layer').bold = True
app_items = [
    'Frontend: React/Next.js app hosted on S3 + CloudFront',
    'API Gateway: RESTful API entry point with request validation and throttling',
    'Lambda Functions: Serverless backend logic (authentication, catalog browsing, order processing)',
    'EC2 Auto Scaling Group: Microservices for complex business logic (2x t3.medium instances)',
    'Application Load Balancer: Distributes traffic to EC2 instances'
]
for item in app_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('3. Data Layer').bold = True
data_items = [
    'RDS PostgreSQL Multi-AZ: User accounts, game catalog, orders, license keys',
    'S3 Buckets: Game installation files, promotional assets, static website hosting',
    'SQS: Message queue for asynchronous order processing and license generation',
    'SNS: Email notifications for order confirmations and promotions'
]
for item in data_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('4. Networking Layer').bold = True
net_items = [
    'VPC: Isolated network environment (10.0.0.0/16)',
    'Public Subnets: ALB, NAT Gateway for internet access',
    'Private Subnets: EC2 instances, RDS database (no direct internet exposure)',
    'Internet Gateway: External access for public-facing resources'
]
for item in net_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('5. CI/CD Layer').bold = True
cicd_items = [
    'GitLab: Source code repository with version control',
    'CodePipeline: Orchestrates automated build and deployment workflow',
    'CodeBuild: Compiles code, runs tests, creates deployment artifacts',
    'EC2 Deployment: Automated deployment to instances via CodeDeploy'
]
for item in cicd_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('6. Security & Monitoring Layer').bold = True
sec_items = [
    'Cognito: User authentication and authorization with MFA support',
    'IAM: Fine-grained access control with least privilege policies',
    'Secrets Manager: Secure storage for database credentials with automatic rotation',
    'CloudWatch: Metrics collection, logs aggregation, and alarms',
    'CloudTrail: API audit logging for compliance'
]
for item in sec_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('User Flow:').bold = True
doc.add_paragraph(
    'User accesses site → Login via Cognito → Browse games (API Gateway → Lambda → RDS) → '
    'Add to cart → Checkout → Payment processing → License generation (SQS) → '
    'Email notification (SNS) → Secure download from S3'
)

doc.add_paragraph('Architecture diagram reference: /images/2-Proposal/proposal.jpg')

# 2.2 TECHNICAL PLAN
doc.add_heading('2.2  TECHNICAL PLAN', 2)
doc.add_paragraph(
    'The partner will implement the architecture using Infrastructure as Code (IaC) via AWS CloudFormation '
    'to ensure consistent, reproducible, and version-controlled deployments.'
)

doc.add_paragraph().add_run('CI/CD Pipeline Stages:').bold = True
pipeline_stages = [
    'Source: GitLab webhook triggers CodePipeline on commit',
    'Build: CodeBuild compiles React frontend and Node.js backend',
    'Test: Automated unit tests, integration tests, and security scans',
    'Deploy: Frontend to S3, Backend to EC2 via CodeDeploy with blue/green deployment'
]
for stage in pipeline_stages:
    doc.add_paragraph(stage, style='List Bullet')

doc.add_paragraph().add_run('Additional Technical Configurations:').bold = True
config_items = [
    'IAM roles restricted following least-privilege model (separate roles for Lambda, EC2, RDS access)',
    'RDS automated backups with 7-day retention and point-in-time recovery',
    'CloudWatch alarms for CPU >80%, memory >85%, database connections >80%, API 5xx errors',
    'S3 versioning enabled for rollback capability',
    'Cognito user pool with strong password policies (min 8 chars, uppercase, lowercase, numbers, symbols)',
    'Lambda provisioned concurrency to eliminate cold starts during peak hours'
]
for config in config_items:
    doc.add_paragraph(config, style='List Bullet')

# 2.3 PROJECT PLAN
doc.add_heading('2.3  PROJECT PLAN', 2)
doc.add_paragraph(
    'The project follows an Agile Scrum framework across 6 months (24 weeks), '
    'organized into 12 sprints of 2 weeks each.'
)

doc.add_paragraph().add_run('Team Structure:').bold = True
team_roles = [
    'Project Manager: Overall coordination, stakeholder communication, risk management',
    'Cloud Solutions Architect: Design AWS infrastructure, create IaC templates, technical leadership',
    'Backend Developers (2): Build Lambda functions, EC2 microservices, API Gateway integration',
    'Frontend Developer: Develop React/Next.js application, Cognito integration',
    'DevOps Engineer: Set up CI/CD pipeline, monitoring, infrastructure automation',
    'QA Engineer: Functional testing, performance testing, security testing'
]
for role in team_roles:
    doc.add_paragraph(role, style='List Bullet')

doc.add_paragraph().add_run('Communication Cadence:').bold = True
comm_items = [
    'Daily standup: 15-minute team sync (blockers, progress, plans)',
    'Weekly sprint planning: Review backlog, assign tasks, estimate effort',
    'Bi-weekly sprint review: Demo completed features to stakeholders',
    'Bi-weekly retrospective: Team improvement discussions',
    'Monthly technical report: Progress update, metrics, risks to customer',
    'Slack/Teams: Real-time collaboration and quick questions'
]
for comm in comm_items:
    doc.add_paragraph(comm, style='List Bullet')

# 2.4 SECURITY CONSIDERATIONS
doc.add_heading('2.4  SECURITY CONSIDERATIONS', 2)

doc.add_paragraph().add_run('Access Control:').bold = True
access_items = [
    'Cognito user pool with email verification and optional MFA (TOTP-based)',
    'IAM roles with least privilege principle',
    'API Gateway request validation, throttling (1000 req/sec), and API keys for partners'
]
for item in access_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Infrastructure Security:').bold = True
infra_sec = [
    'EC2 instances in private subnets (no direct internet access)',
    'Security Groups with minimal required ports (ALB: 443/80, EC2: 8080, RDS: 5432)',
    'CloudFront with HTTPS-only, AWS Shield Standard for DDoS protection'
]
for item in infra_sec:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Data Protection:').bold = True
data_prot = [
    'RDS encryption at rest using AWS KMS (AES-256)',
    'S3 bucket encryption (SSE-S3) and versioning enabled',
    'SSL/TLS for all data in transit (ACM-issued certificates)',
    'Secrets Manager for database credentials with automatic 30-day rotation'
]
for item in data_prot:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Detection & Response:').bold = True
detect_items = [
    'CloudTrail for API audit logging (all actions tracked, 90-day retention)',
    'CloudWatch alarms for failed login attempts, unusual traffic patterns, high CPU/memory',
    'SNS notifications for critical security events',
    'Regular security audits (quarterly) and penetration testing (before launch)'
]
for item in detect_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph().add_run('Incident Management:').bold = True
incident_items = [
    'Automated backup and restore procedures documented in runbooks',
    'RDS automated backups with 7-day retention and point-in-time recovery',
    'Disaster recovery plan: RTO < 4 hours, RPO < 1 hour',
    'Incident response runbook for common scenarios (DB failover, DDoS, data breach)'
]
for item in incident_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 3: ACTIVITIES AND DELIVERABLES
doc.add_heading('3  ACTIVITIES AND DELIVERABLES', 1)

doc.add_heading('3.1  ACTIVITIES AND DELIVERABLES', 2)

table = doc.add_table(rows=8, cols=5)
table.style = 'Light Grid Accent 1'

# Headers
headers = ['Phase', 'Timeline', 'Activities', 'Deliverables', 'Man-days']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Data rows
data = [
    ('Assessment', 'Week 1-2', 'Requirements gathering, architecture design, ERD modeling, cost estimation', 'Architecture document, project plan, ERD diagrams', '10'),
    ('Infrastructure', 'Week 3-6', 'VPC, subnets, IGW, NAT, RDS, S3, Cognito setup', 'AWS environment ready, networking configured', '15'),
    ('Backend Dev', 'Week 7-12', 'Lambda APIs, EC2 microservices, API Gateway integration', 'Working APIs for auth, catalog, orders', '25'),
    ('Frontend Dev', 'Week 11-14', 'React app, Cognito integration, S3 + CloudFront deployment', 'User-facing web application', '18'),
    ('Integration', 'Week 15-18', 'Payment gateway, admin dashboard, CI/CD setup', 'End-to-end workflow functional', '20'),
    ('Testing', 'Week 19-22', 'Load testing, security audits, UAT, bug fixes', 'Test reports, security audit', '15'),
    ('Launch', 'Week 23-24', 'Beta release, monitoring, production deployment, handover', 'Production system, documentation', '7')
]

for i, row_data in enumerate(data, start=1):
    for j, value in enumerate(row_data):
        table.rows[i].cells[j].text = value

# 3.2 OUT OF SCOPE
doc.add_heading('3.2  OUT OF SCOPE', 2)
doc.add_paragraph('The following items are excluded from the initial project scope:')
out_of_scope = [
    'AI/ML recommendation engine for personalized game suggestions',
    'Mobile application development (iOS/Android native apps)',
    'Integration with external inventory management or ERP systems',
    'Multi-region failover and disaster recovery setup',
    'Advanced analytics and business intelligence dashboards',
    'Social media login integration (Google, Facebook)',
    'Live chat customer support system',
    'Advanced fraud detection beyond basic payment gateway features',
    'Game streaming or cloud gaming capabilities',
    'Marketplace features (user-to-user game trading)'
]
for item in out_of_scope:
    doc.add_paragraph(item, style='List Bullet')

# 3.3 PATH TO PRODUCTION
doc.add_heading('3.3  PATH TO PRODUCTION', 2)
doc.add_paragraph('The initial release (Phase 1 / MVP) includes core functionalities:')
mvp_features = [
    'User registration and authentication via Cognito',
    'Game catalog browsing with search and filtering',
    'Shopping cart and checkout workflow',
    'Payment processing via Stripe/PayPal integration',
    'Automated license key generation and delivery',
    'Secure game file downloads from S3',
    'Admin dashboard for inventory management',
    'Basic reporting and analytics'
]
for feature in mvp_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('To achieve production-ready status, the following enhancements are planned:')
prod_ready = [
    'Enhanced error handling and user-friendly error messages',
    'API rate limiting and request throttling via API Gateway',
    'Database query optimization and indexing',
    'Comprehensive audit trail dashboards in CloudWatch',
    'Penetration testing by third-party security firm',
    'Load testing to validate 10x traffic spike handling',
    'Customer acceptance testing (UAT) with beta users'
]
for item in prod_ready:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 4: AWS COST BREAKDOWN
doc.add_heading('4  EXPECTED AWS COST BREAKDOWN BY SERVICES', 1)
doc.add_paragraph('Estimated monthly cost for 10,000 users and 1,000 orders/month: ~$228')

cost_table = doc.add_table(rows=10, cols=2)
cost_table.style = 'Light Grid Accent 1'

cost_headers = ['AWS Service', 'Monthly Cost (USD)']
for i, header in enumerate(cost_headers):
    cell = cost_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

cost_data = [
    ('CloudFront (10M requests, 50GB transfer)', '$8'),
    ('S3 (100GB storage + requests)', '$3'),
    ('API Gateway (1M requests)', '$3.50'),
    ('Lambda (5M invocations, 512MB)', '$17.50'),
    ('EC2 (2x t3.medium instances)', '$60'),
    ('RDS PostgreSQL (db.t3.small Multi-AZ)', '$50'),
    ('Application Load Balancer', '$22'),
    ('Cognito (10K users)', '$28'),
    ('SQS + SNS + CloudWatch + Other', '$36')
]

for i, (service, cost) in enumerate(cost_data, start=1):
    cost_table.rows[i].cells[0].text = service
    cost_table.rows[i].cells[1].text = cost

doc.add_paragraph()
doc.add_paragraph('Total Monthly Cost: ~$228')
doc.add_paragraph('Estimated Annual Cost: ~$2,736')

doc.add_paragraph().add_run('Scaling Projections:').bold = True
doc.add_paragraph('• 50,000 users: ~$650/month')
doc.add_paragraph('• 100,000 users: ~$1,200/month')

doc.add_paragraph().add_run('Assumptions:').bold = True
cost_assumptions = [
    'Traffic under 10,000 users/month initially',
    'AWS Free Tier credits applied where eligible',
    'Data transfer within same AWS region'
]
for assumption in cost_assumptions:
    doc.add_paragraph(assumption, style='List Bullet')

doc.add_page_break()

# Section 5: TEAM
doc.add_heading('5  TEAM', 1)

doc.add_paragraph().add_run('Partner Executive Sponsor').bold = True
exec_table = doc.add_table(rows=2, cols=4)
exec_table.style = 'Light Grid Accent 1'

exec_headers = ['Name', 'Title', 'Description', 'Email / Contact Info']
for i, header in enumerate(exec_headers):
    cell = exec_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

exec_table.rows[1].cells[0].text = 'Nguyen Van A'
exec_table.rows[1].cells[1].text = 'Director of Cloud Services'
exec_table.rows[1].cells[2].text = 'Overall project accountability and strategic direction'
exec_table.rows[1].cells[3].text = 'nguyenvana@devteria.com'

doc.add_paragraph()
doc.add_paragraph().add_run('Project Stakeholders').bold = True
stakeholder_table = doc.add_table(rows=4, cols=4)
stakeholder_table.style = 'Light Grid Accent 1'

stakeholder_headers = ['Name', 'Title', 'Stakeholder for', 'Email / Contact Info']
for i, header in enumerate(stakeholder_headers):
    cell = stakeholder_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

stakeholders = [
    ('Tran Thi B', 'Product Manager', 'Business requirements, UX decisions', 'tranthib@devteria.com'),
    ('Le Van C', 'Senior QA Engineer', 'Testing and quality assurance', 'levanc@devteria.com'),
    ('Pham Thi D', 'IT Operations Manager', 'Infrastructure and deployment', 'phamthid@devteria.com')
]

for i, (name, title, stakeholder_for, email) in enumerate(stakeholders, start=1):
    stakeholder_table.rows[i].cells[0].text = name
    stakeholder_table.rows[i].cells[1].text = title
    stakeholder_table.rows[i].cells[2].text = stakeholder_for
    stakeholder_table.rows[i].cells[3].text = email

doc.add_paragraph()
doc.add_paragraph().add_run('Partner Project Team').bold = True
team_table = doc.add_table(rows=7, cols=4)
team_table.style = 'Light Grid Accent 1'

team_headers = ['Name', 'Title', 'Role', 'Email / Contact Info']
for i, header in enumerate(team_headers):
    cell = team_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

team_members = [
    ('Hoang Van E', 'Project Manager', 'Project coordination', 'hoangvane@partner.com'),
    ('Nguyen Thi F', 'Cloud Solutions Architect', 'Architecture design, technical lead', 'nguyenthif@partner.com'),
    ('Tran Van G', 'Senior Backend Developer', 'Lambda and EC2 development', 'tranvang@partner.com'),
    ('Pham Thi H', 'Frontend Developer', 'React application development', 'phamthih@partner.com'),
    ('Le Van I', 'DevOps Engineer', 'CI/CD and infrastructure automation', 'levani@partner.com'),
    ('Nguyen Van J', 'QA Engineer', 'Testing and quality assurance', 'nguyenvanj@partner.com')
]

for i, (name, title, role, email) in enumerate(team_members, start=1):
    team_table.rows[i].cells[0].text = name
    team_table.rows[i].cells[1].text = title
    team_table.rows[i].cells[2].text = role
    team_table.rows[i].cells[3].text = email

doc.add_page_break()

# Section 6: RESOURCES & COST ESTIMATES
doc.add_heading('6  RESOURCES & COST ESTIMATES', 1)

doc.add_paragraph().add_run('Resource Rates:').bold = True
rate_table = doc.add_table(rows=4, cols=2)
rate_table.style = 'Light Grid Accent 1'

rate_headers = ['Resource', 'Rate (USD) / Hour']
for i, header in enumerate(rate_headers):
    cell = rate_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

rates = [
    ('Solution Architects (1 headcount)', '$80'),
    ('Engineers (4 headcount)', '$60'),
    ('QA Engineer (1 headcount)', '$55')
]

for i, (resource, rate) in enumerate(rates, start=1):
    rate_table.rows[i].cells[0].text = resource
    rate_table.rows[i].cells[1].text = rate

doc.add_paragraph()
doc.add_paragraph().add_run('Project Phase Hours Breakdown:').bold = True

hours_table = doc.add_table(rows=9, cols=5)
hours_table.style = 'Light Grid Accent 1'

hours_headers = ['Project Phase', 'Solution Architects', 'Engineers', 'QA Engineer', 'Total Hours']
for i, header in enumerate(hours_headers):
    cell = hours_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

hours_data = [
    ('Assessment', '80', '0', '0', '80'),
    ('Infrastructure', '120', '0', '0', '120'),
    ('Backend Dev', '40', '200', '0', '240'),
    ('Frontend Dev', '0', '144', '0', '144'),
    ('Integration', '40', '160', '0', '200'),
    ('Testing', '20', '0', '120', '140'),
    ('Launch', '16', '40', '0', '56'),
    ('Total Hours', '316', '544', '120', '980')
]

for i, row_data in enumerate(hours_data, start=1):
    for j, value in enumerate(row_data):
        cell = hours_table.rows[i].cells[j]
        cell.text = value
        if i == len(hours_data):  # Bold total row
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

doc.add_paragraph()
doc.add_paragraph().add_run('Total Cost Calculation:').bold = True
doc.add_paragraph('• Solution Architects: 316 hours × $80/hour = $25,280')
doc.add_paragraph('• Engineers: 544 hours × $60/hour = $32,640')
doc.add_paragraph('• QA Engineer: 120 hours × $55/hour = $6,600')
doc.add_paragraph().add_run('Total Project Cost: $64,520').bold = True

doc.add_paragraph()
doc.add_paragraph().add_run('Cost Contribution Distribution:').bold = True

contrib_table = doc.add_table(rows=4, cols=3)
contrib_table.style = 'Light Grid Accent 1'

contrib_headers = ['Party', 'Contribution (USD)', '% of Total']
for i, header in enumerate(contrib_headers):
    cell = contrib_table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

contrib_data = [
    ('Customer', '$32,260', '50%'),
    ('Partner', '$25,808', '40%'),
    ('AWS', '$6,452', '10%')
]

for i, (party, amount, percent) in enumerate(contrib_data, start=1):
    contrib_table.rows[i].cells[0].text = party
    contrib_table.rows[i].cells[1].text = amount
    contrib_table.rows[i].cells[2].text = percent

doc.add_page_break()

# Section 7: ACCEPTANCE
doc.add_heading('7  ACCEPTANCE', 1)

doc.add_paragraph(
    'Upon completion of each deliverable listed in Section 3.1, the partner will submit deliverables '
    'along with a formal Acceptance Form for customer review.'
)

doc.add_paragraph().add_run('Acceptance Process:').bold = True
acceptance_process = [
    'Partner completes deliverable and internal quality checks',
    'Partner submits deliverable with Acceptance Form to customer',
    'Customer has 8 business days from submission date to review'
]
for step in acceptance_process:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph().add_run('Customer Review Options:').bold = True
review_options = [
    'Approve: Deliverable meets acceptance criteria, project proceeds to next phase',
    'Reject: Deliverable does not meet criteria, partner receives detailed feedback specifying required changes'
]
for option in review_options:
    doc.add_paragraph(option, style='List Bullet')

doc.add_paragraph().add_run('Rejection Handling:').bold = True
rejection_steps = [
    'Partner corrects deliverable according to feedback',
    'Partner resubmits for another review cycle (another 8 business days)',
    'Process repeats until approval or escalation'
]
for step in rejection_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph().add_run('Auto-Acceptance:').bold = True
doc.add_paragraph(
    'If customer does not respond within 8-business-day acceptance period, deliverable is automatically accepted '
    'and project proceeds to next phase.'
)

doc.add_paragraph().add_run('Documentation:').bold = True
doc_items = [
    'All acceptances (explicit or auto) will be documented and signed by both parties',
    'Maintains clear audit trail of project progress and completion milestones',
    'Signatures required from Customer Executive Sponsor and Partner Project Manager'
]
for item in doc_items:
    doc.add_paragraph(item, style='List Bullet')

# Save document
doc.save('c:/Users/LENOVO/Documents/GitHub/Baovc_AWS_InternReport/AWS_Project_Plan_Devteria.docx')
print('✅ Word document created successfully: AWS_Project_Plan_Devteria.docx')
